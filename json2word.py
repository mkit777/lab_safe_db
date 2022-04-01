from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX, MSO_COLOR_TYPE

import json

if __name__ == "__main__":
    f = open('实验室安全教育题库.json')
    data = json.load(f)

    doc = Document()
    doc.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体' # 添加字体样式-Song
    doc.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    doc.add_heading('实验室安全教育题库（全）', 0)
   
    TYPE_MAP = {
        'SingleSelection':'单选',
        'Judgement':'判断',
        'MultipleSelection':'多选',
        'TextEntry': '填空',
        'Composition': '简答'
    }

    for idx, item in enumerate(data):
        print(f'Progress {idx+1}/{len(data)}')
        typeStr = TYPE_MAP[item['type']]

        no = int(item['rel'])
        doc.add_paragraph().add_run(f'{idx+1}.{item["title"]}【{typeStr}】', style='Song').bold = True

        answer = item["answer"]
        
        if typeStr not in ['填空', '简答']:
            for selection in item['selections']:
                if selection[:1] in answer:
                    p = doc.add_paragraph().add_run(selection, style='Song')
                    p.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
                    p.bold = True
                else:
                    doc.add_paragraph().add_run(selection, style='Song')
           
        doc.add_paragraph().add_run(f'答案：{answer}', style='Song')

    doc.save('实验室安全教育题库.docx')